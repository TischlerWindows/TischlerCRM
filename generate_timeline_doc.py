"""
Generates "TischlerCRM Project Timeline.docx" — run with:
    pip install python-docx
    python generate_timeline_doc.py
Output written next to this script.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1B, 0x2A, 0x4A)
RED = RGBColor(0xC0, 0x1F, 0x2E)
GREEN = RGBColor(0x1E, 0x7A, 0x34)
AMBER = RGBColor(0xB8, 0x86, 0x00)
GRAY = RGBColor(0x55, 0x55, 0x55)

STATUS_COLOR = {
    'DONE': GREEN,
    'DONE*': GREEN,
    'PARTIAL': AMBER,
    'IN PROGRESS': AMBER,
    'NOT DONE': RED,
}


def set_cell_shading(cell, hex_color):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def add_heading(doc, text, size=20, color=NAVY, space_before=18, space_after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p


def add_status_table(doc, rows):
    """rows: list of (item, status, note)"""
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = 'Item'
    hdr[1].text = 'Status'
    hdr[2].text = 'Notes'
    for cell in hdr:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, '1B2A4A')

    widths = [Inches(2.1), Inches(1.1), Inches(3.6)]
    for item, status, note in rows:
        row = table.add_row().cells
        row[0].text = item
        row[1].text = status
        row[2].text = note
        color = STATUS_COLOR.get(status, GRAY)
        for p in row[1].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = color
        for i, w in enumerate(widths):
            row[i].width = w
    for i, w in enumerate(widths):
        hdr[i].width = w
    return table


def main():
    doc = Document()

    # Base style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('TischlerCRM — Project Timeline')
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = NAVY

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run('Status as of July 30, 2026')
    run.italic = True
    run.font.size = Pt(12)
    run.font.color.rgb = GRAY

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run(
        'Original plan assumed AWS (Lambda/RDS/Cognito/Terraform/SES). '
        'Actual build uses Railway + Fastify + custom JWT auth — simpler, and in places well beyond the original scope.'
    )
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY
    run.italic = True

    # ── Phase 1 ──────────────────────────────────────────────
    add_heading(doc, 'Phase 1 — Foundations  ✅ Complete')
    add_status_table(doc, [
        ('Repo / monorepo / CI', 'DONE', 'pnpm workspaces (web, api, db, types, storage, widgets, triggers, controllers) + GitHub Actions'),
        ('Infra', 'DONE*', 'Railway instead of AWS/Terraform/Lambda — no IaC needed, Railway manages it'),
        ('Database / ORM', 'DONE*', 'Prisma + Postgres, but built as a Salesforce-style custom-object metadata layer, not fixed tables — exceeds original scope'),
        ('Auth', 'DONE*', 'Custom JWT ({sub, role, exp}) instead of Cognito Hosted UI'),
        ('Monitoring', 'PARTIAL', 'No CloudWatch (N/A off AWS); Sentry / error-tracking status unconfirmed'),
    ])

    # ── Phase 2 ──────────────────────────────────────────────
    add_heading(doc, 'Phase 2 — Core CRM Features  ✅ Mostly complete')
    add_status_table(doc, [
        ('Core entity CRUD', 'DONE*', 'Every custom object gets full CRUD + a visual page-layout builder, not just Accounts/Contacts/Opportunities'),
        ('Activities / Notes', 'PARTIAL', 'Likely via the generic Record system rather than a dedicated Activities table — confirm timeline/notes UI on record pages'),
        ('Dropbox integration', 'DONE*', 'Native OAuth + file browser (packages/storage) — Zapier phase skipped entirely'),
        ('Pipeline dashboards / reporting', 'DONE', 'Dashboard API + reference docs exist in repo'),
        ('Global search', 'DONE', 'Search implementation present'),
    ])

    # ── Phase 3 ──────────────────────────────────────────────
    add_heading(doc, 'Phase 3 — Security, Permissions & Scale  ⚠️ Partial')
    add_status_table(doc, [
        ('RBAC', 'PARTIAL', 'Simple Admin/User role flag today, not the full Admin/Manager/User/Viewer tiers'),
        ('Validation / error handling', 'DONE', 'Zod on every request body is a hard repo convention'),
        ('Performance', 'PARTIAL', 'limit/offset clamping is convention; no confirmed load-time benchmarking'),
        ('Notifications / Automations', 'DONE*', 'Full Triggers/Controllers/Widgets engine (opt-out model) — well beyond the "simple workflow builder" originally scoped'),
        ('Email sending', 'IN PROGRESS', 'Multi-week saga: Outlook SMTP blocked, Graph OAuth 403, Resend/SendGrid need domain, Zapier hit dead ends. Mailjet setup underway, not yet confirmed live'),
        ('Backups / Audit', 'DONE*', 'AuditLog + LoginEvent models, full JSON snapshot export/restore — on-demand, not scheduled daily'),
    ])

    # ── Phase 4 ──────────────────────────────────────────────
    add_heading(doc, 'Phase 4 — UX Polish, Testing & Rollout  🔄 In progress')
    add_status_table(doc, [
        ('UX polish', 'IN PROGRESS', 'Brand guide exists; page-editor bulk-formatting feature just shipped'),
        ('QA / E2E testing', 'NOT DONE', 'Jest covers apps/web only; no Cypress/Playwright suite, no backend tests — explicit repo decision, not an oversight'),
        ('Documentation', 'DONE*', 'Extensive internal technical docs already in repo (architecture, dashboard API, migrations, deployment, integrations)'),
        ('User training material', 'NOT DONE', 'Admin/user guides, Loom videos — not yet produced'),
        ('Deployment / rollout', 'DONE', 'Live in production at tischlercrm.up.railway.app with real daily use'),
        ('Stabilization', 'IN PROGRESS', 'Continuous — this week alone: formula-field bug, page-editor duplicate-field bug, lookup-cache bug, static-URL-field bug, debug-toast cleanup'),
    ])

    # ── Next Phases ──────────────────────────────────────────
    add_heading(doc, 'Next Phases — What\u2019s Ahead', color=RED)
    add_status_table(doc, [
        ('Finish Project List Report', 'IN PROGRESS', 'Complete remaining Project object fields/layout to match the Tischler master project-tracking sheet 1:1'),
        ('Finish Service structure', 'IN PROGRESS', 'Finalize Service object schema, layouts, and workflow to match real support/service processes'),
        ('Finish Email configuration', 'IN PROGRESS', 'Resolve outbound email via Mailjet (single-sender verification, in progress) OR set up a real owned domain for Resend/SendGrid with proper SPF/DKIM — either closes out the entire email saga permanently'),
        ('RBAC — full role tiers', 'NOT STARTED', 'Expand from Admin/User to Admin/Manager/User/Viewer with granular permission checks'),
        ('Automated testing', 'NOT STARTED', 'No E2E suite currently planned; revisit if regressions become frequent'),
        ('User-facing documentation', 'NOT STARTED', 'Admin & user guides, short training videos for team onboarding'),
        ('Scheduled/automated backups', 'NOT STARTED', 'Current backup system is on-demand only; consider a scheduled job'),
    ])

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run('Generated for internal planning purposes — TischlerCRM')
    run.italic = True
    run.font.size = Pt(8)
    run.font.color.rgb = GRAY

    doc.save('TischlerCRM Project Timeline.docx')
    print('Saved: TischlerCRM Project Timeline.docx')


if __name__ == '__main__':
    main()
